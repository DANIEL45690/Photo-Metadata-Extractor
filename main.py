import os
import sys
import json
import csv
import sqlite3
import hashlib
import time
import uuid
import shutil
import zipfile
import tarfile
import re
import math
import socket
import urllib.parse
import mimetypes
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
import warnings
warnings.filterwarnings('ignore')

import requests
import cv2
import numpy as np
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
import folium
from folium.plugins import HeatMap, MarkerCluster
import webbrowser

DB_PATH = "scanner_data.db"
REPORT_PATH = "scanner_reports"
JSON_PATH = "json_exports"
GPS_PATH = "gps_locations"
os.makedirs(REPORT_PATH, exist_ok=True)
os.makedirs(JSON_PATH, exist_ok=True)
os.makedirs(GPS_PATH, exist_ok=True)

print(f"[SYSTEM] Database: {os.path.abspath(DB_PATH)}")
print(f"[SYSTEM] Reports: {os.path.abspath(REPORT_PATH)}")
print(f"[SYSTEM] JSON exports: {os.path.abspath(JSON_PATH)}")
print(f"[SYSTEM] GPS maps: {os.path.abspath(GPS_PATH)}")

class DatabaseManager:
    def __init__(self):
        self.db_path = DB_PATH
        self.conn = sqlite3.connect(self.db_path)
        self.conn.row_factory = sqlite3.Row
        self.cursor = self.conn.cursor()
        self.init_tables()
        print(f"[DB] Database initialized at: {self.db_path}")

    def init_tables(self):
        self.cursor.execute('DROP TABLE IF EXISTS scans')
        self.cursor.execute('DROP TABLE IF EXISTS files')
        self.cursor.execute('DROP TABLE IF EXISTS gps_data')
        self.cursor.execute('DROP TABLE IF EXISTS ip_data')
        self.cursor.execute('DROP TABLE IF EXISTS documents')
        self.cursor.execute('DROP TABLE IF EXISTS audio_data')
        self.cursor.execute('DROP TABLE IF EXISTS video_data')
        self.cursor.execute('DROP TABLE IF EXISTS archives')
        self.cursor.execute('DROP TABLE IF EXISTS exif_data')
        self.cursor.execute('DROP TABLE IF EXISTS faces_data')
        self.cursor.execute('DROP TABLE IF EXISTS ocr_data')
        self.cursor.execute('DROP TABLE IF EXISTS stego_data')
        self.cursor.execute('DROP TABLE IF EXISTS network_info')
        self.cursor.execute('DROP TABLE IF EXISTS system_info')

        self.cursor.execute('''CREATE TABLE scans (
            scan_id TEXT PRIMARY KEY,
            start_time TEXT,
            end_time TEXT,
            scan_path TEXT,
            total_files INTEGER,
            total_gps INTEGER,
            total_documents INTEGER,
            total_audio INTEGER,
            total_video INTEGER,
            total_archives INTEGER,
            total_stego INTEGER,
            total_faces INTEGER,
            duration_seconds REAL,
            scan_type TEXT,
            computer_name TEXT,
            json_export_path TEXT
        )''')

        self.cursor.execute('''CREATE TABLE files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            file_name TEXT,
            file_size INTEGER,
            sha256 TEXT,
            md5 TEXT,
            mime_type TEXT,
            extension TEXT,
            created_time TEXT,
            modified_time TEXT
        )''')

        self.cursor.execute('''CREATE TABLE gps_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            latitude REAL,
            longitude REAL,
            altitude REAL,
            accuracy REAL,
            direction REAL,
            address TEXT,
            country TEXT,
            city TEXT,
            street TEXT,
            postcode TEXT,
            google_maps_url TEXT,
            osm_url TEXT
        )''')

        self.cursor.execute('''CREATE TABLE ip_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            ip_address TEXT,
            latitude REAL,
            longitude REAL,
            city TEXT,
            country TEXT,
            region TEXT,
            isp TEXT,
            organization TEXT,
            asn TEXT,
            timezone TEXT,
            source_name TEXT
        )''')

        self.cursor.execute('''CREATE TABLE documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            page_count INTEGER,
            word_count INTEGER,
            character_count INTEGER,
            author TEXT,
            title TEXT,
            creator TEXT,
            creation_date TEXT,
            full_text TEXT
        )''')

        self.cursor.execute('''CREATE TABLE audio_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            duration_seconds REAL,
            duration_formatted TEXT,
            bitrate INTEGER,
            sample_rate INTEGER,
            channels INTEGER,
            artist TEXT,
            album TEXT,
            title TEXT,
            genre TEXT,
            year INTEGER
        )''')

        self.cursor.execute('''CREATE TABLE video_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            duration_seconds REAL,
            duration_formatted TEXT,
            width INTEGER,
            height INTEGER,
            fps REAL,
            bitrate INTEGER,
            video_codec TEXT,
            frame_count INTEGER
        )''')

        self.cursor.execute('''CREATE TABLE archives (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            archive_type TEXT,
            file_count INTEGER,
            total_size INTEGER,
            files_list TEXT
        )''')

        self.cursor.execute('''CREATE TABLE exif_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            camera_make TEXT,
            camera_model TEXT,
            lens_model TEXT,
            focal_length TEXT,
            aperture TEXT,
            shutter_speed TEXT,
            iso INTEGER,
            date_taken TEXT,
            image_width INTEGER,
            image_height INTEGER,
            orientation INTEGER
        )''')

        self.cursor.execute('''CREATE TABLE faces_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            face_count INTEGER,
            faces_json TEXT
        )''')

        self.cursor.execute('''CREATE TABLE ocr_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            extracted_text TEXT,
            word_count INTEGER,
            confidence REAL
        )''')

        self.cursor.execute('''CREATE TABLE stego_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            file_path TEXT,
            has_hidden_data INTEGER,
            detection_method TEXT,
            confidence REAL
        )''')

        self.cursor.execute('''CREATE TABLE network_info (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            interface_name TEXT,
            mac_address TEXT,
            ipv4 TEXT,
            gateway TEXT,
            external_ip TEXT
        )''')

        self.cursor.execute('''CREATE TABLE system_info (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            scan_id TEXT,
            hostname TEXT,
            os_name TEXT,
            os_version TEXT,
            cpu_model TEXT,
            cpu_cores INTEGER,
            ram_total_gb REAL,
            disk_total_gb REAL,
            python_version TEXT
        )''')

        self.conn.commit()

    def save_scan(self, data):
        self.cursor.execute('INSERT INTO scans VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('start_time'), data.get('end_time'),
             data.get('scan_path'), data.get('total_files'), data.get('total_gps'),
             data.get('total_documents'), data.get('total_audio'), data.get('total_video'),
             data.get('total_archives'), data.get('total_stego'), data.get('total_faces'),
             data.get('duration_seconds'), data.get('scan_type'), data.get('computer_name'),
             data.get('json_export_path')))
        self.conn.commit()
        print(f"[DB] Scan saved: {data.get('scan_id')}")

    def save_file(self, data):
        self.cursor.execute('INSERT INTO files (scan_id, file_path, file_name, file_size, sha256, md5, mime_type, extension, created_time, modified_time) VALUES (?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('file_name'),
             data.get('file_size'), data.get('sha256'), data.get('md5'),
             data.get('mime_type'), data.get('extension'), data.get('created_time'),
             data.get('modified_time')))
        self.conn.commit()

    def save_gps(self, data):
        self.cursor.execute('INSERT INTO gps_data (scan_id, file_path, latitude, longitude, altitude, accuracy, direction, address, country, city, street, postcode, google_maps_url, osm_url) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('latitude'),
             data.get('longitude'), data.get('altitude'), data.get('accuracy'),
             data.get('direction'), data.get('address'), data.get('country'),
             data.get('city'), data.get('street'), data.get('postcode'),
             data.get('google_maps_url'), data.get('osm_url')))
        self.conn.commit()
        print(f"[DB] GPS saved: {data.get('latitude')}, {data.get('longitude')}")

    def save_ip(self, data):
        self.cursor.execute('INSERT INTO ip_data (scan_id, ip_address, latitude, longitude, city, country, region, isp, organization, asn, timezone, source_name) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('ip_address'), data.get('latitude'),
             data.get('longitude'), data.get('city'), data.get('country'),
             data.get('region'), data.get('isp'), data.get('organization'),
             data.get('asn'), data.get('timezone'), data.get('source_name')))
        self.conn.commit()

    def save_document(self, data):
        self.cursor.execute('INSERT INTO documents (scan_id, file_path, page_count, word_count, character_count, author, title, creator, creation_date, full_text) VALUES (?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('page_count'),
             data.get('word_count'), data.get('character_count'), data.get('author'),
             data.get('title'), data.get('creator'), data.get('creation_date'),
             data.get('full_text')))
        self.conn.commit()

    def save_audio(self, data):
        self.cursor.execute('INSERT INTO audio_data (scan_id, file_path, duration_seconds, duration_formatted, bitrate, sample_rate, channels, artist, album, title, genre, year) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('duration_seconds'),
             data.get('duration_formatted'), data.get('bitrate'), data.get('sample_rate'),
             data.get('channels'), data.get('artist'), data.get('album'),
             data.get('title'), data.get('genre'), data.get('year')))
        self.conn.commit()

    def save_video(self, data):
        self.cursor.execute('INSERT INTO video_data (scan_id, file_path, duration_seconds, duration_formatted, width, height, fps, bitrate, video_codec, frame_count) VALUES (?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('duration_seconds'),
             data.get('duration_formatted'), data.get('width'), data.get('height'),
             data.get('fps'), data.get('bitrate'), data.get('video_codec'),
             data.get('frame_count')))
        self.conn.commit()

    def save_archive(self, data):
        self.cursor.execute('INSERT INTO archives (scan_id, file_path, archive_type, file_count, total_size, files_list) VALUES (?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('archive_type'),
             data.get('file_count'), data.get('total_size'), data.get('files_list')))
        self.conn.commit()

    def save_exif(self, data):
        self.cursor.execute('INSERT INTO exif_data (scan_id, file_path, camera_make, camera_model, lens_model, focal_length, aperture, shutter_speed, iso, date_taken, image_width, image_height, orientation) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('camera_make'),
             data.get('camera_model'), data.get('lens_model'), data.get('focal_length'),
             data.get('aperture'), data.get('shutter_speed'), data.get('iso'),
             data.get('date_taken'), data.get('image_width'), data.get('image_height'),
             data.get('orientation')))
        self.conn.commit()

    def save_faces(self, data):
        self.cursor.execute('INSERT INTO faces_data (scan_id, file_path, face_count, faces_json) VALUES (?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('face_count'),
             data.get('faces_json')))
        self.conn.commit()

    def save_ocr(self, data):
        self.cursor.execute('INSERT INTO ocr_data (scan_id, file_path, extracted_text, word_count, confidence) VALUES (?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('extracted_text'),
             data.get('word_count'), data.get('confidence')))
        self.conn.commit()

    def save_stego(self, data):
        self.cursor.execute('INSERT INTO stego_data (scan_id, file_path, has_hidden_data, detection_method, confidence) VALUES (?,?,?,?,?)',
            (data.get('scan_id'), data.get('file_path'), data.get('has_hidden_data'),
             data.get('detection_method'), data.get('confidence')))
        self.conn.commit()

    def save_network(self, data):
        self.cursor.execute('INSERT INTO network_info (scan_id, interface_name, mac_address, ipv4, gateway, external_ip) VALUES (?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('interface_name'), data.get('mac_address'),
             data.get('ipv4'), data.get('gateway'), data.get('external_ip')))
        self.conn.commit()

    def save_system(self, data):
        self.cursor.execute('INSERT INTO system_info (scan_id, hostname, os_name, os_version, cpu_model, cpu_cores, ram_total_gb, disk_total_gb, python_version) VALUES (?,?,?,?,?,?,?,?,?)',
            (data.get('scan_id'), data.get('hostname'), data.get('os_name'),
             data.get('os_version'), data.get('cpu_model'), data.get('cpu_cores'),
             data.get('ram_total_gb'), data.get('disk_total_gb'), data.get('python_version')))
        self.conn.commit()

    def get_all_gps(self):
        self.cursor.execute('SELECT file_path, latitude, longitude, address, country, city FROM gps_data WHERE latitude IS NOT NULL')
        return self.cursor.fetchall()

    def get_gps_by_scan(self, scan_id):
        self.cursor.execute('SELECT * FROM gps_data WHERE scan_id=?', (scan_id,))
        return self.cursor.fetchall()

    def get_scans(self):
        self.cursor.execute('SELECT * FROM scans ORDER BY start_time DESC')
        return self.cursor.fetchall()

    def get_files_by_scan(self, scan_id):
        self.cursor.execute('SELECT * FROM files WHERE scan_id=?', (scan_id,))
        return self.cursor.fetchall()

    def scan_exists(self, scan_id):
        self.cursor.execute('SELECT scan_id FROM scans WHERE scan_id=?', (scan_id,))
        return self.cursor.fetchone() is not None

    def close(self):
        self.conn.close()

class FileHasher:
    @staticmethod
    def calculate(file_path):
        result = {'sha256': None, 'md5': None}
        try:
            sha256 = hashlib.sha256()
            md5 = hashlib.md5()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(65536), b''):
                    sha256.update(chunk)
                    md5.update(chunk)
            result['sha256'] = sha256.hexdigest()
            result['md5'] = md5.hexdigest()
        except Exception as e:
            print(f"[ERROR] Hash calculation failed: {e}")
        return result

class FileInfoGatherer:
    def __init__(self):
        pass

    def get_info(self, file_path, scan_id):
        try:
            stat = os.stat(file_path)
            hashes = FileHasher.calculate(file_path)
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = 'application/octet-stream'
            return {
                'scan_id': scan_id,
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_size': stat.st_size,
                'sha256': hashes['sha256'],
                'md5': hashes['md5'],
                'mime_type': mime_type,
                'extension': os.path.splitext(file_path)[1].lower(),
                'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat()
            }
        except Exception as e:
            print(f"[ERROR] File info failed: {e}")
            return None

class GPSProcessor:
    @staticmethod
    def dms_to_decimal(dms):
        try:
            if isinstance(dms, tuple) and len(dms) >= 3:
                return float(dms[0]) + float(dms[1])/60 + float(dms[2])/3600
            return 0
        except:
            return 0

    @staticmethod
    def extract(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path}
            img = Image.open(file_path)
            exif = img._getexif()
            if not exif:
                return None

            gps_info = {}
            for k, v in exif.items():
                tag = TAGS.get(k, k)
                if tag == 'GPSInfo':
                    for gk, gv in v.items():
                        gtag = GPSTAGS.get(gk, gk)
                        gps_info[gtag] = gv

            if 'GPSLatitude' not in gps_info or 'GPSLongitude' not in gps_info:
                return None

            lat = GPSProcessor.dms_to_decimal(gps_info['GPSLatitude'])
            lon = GPSProcessor.dms_to_decimal(gps_info['GPSLongitude'])

            if gps_info.get('GPSLatitudeRef') == 'S':
                lat = -lat
            if gps_info.get('GPSLongitudeRef') == 'W':
                lon = -lon

            result['latitude'] = lat
            result['longitude'] = lon
            result['google_maps_url'] = f"https://www.google.com/maps?q={lat},{lon}"
            result['osm_url'] = f"https://www.openstreetmap.org/?mlat={lat}&mlon={lon}"

            if 'GPSAltitude' in gps_info:
                alt = float(gps_info['GPSAltitude'])
                if gps_info.get('GPSAltitudeRef', 0) == 1:
                    alt = -alt
                result['altitude'] = alt
            if 'GPSImgDirection' in gps_info:
                result['direction'] = float(gps_info['GPSImgDirection'])
            if 'GPSDOP' in gps_info:
                result['accuracy'] = float(gps_info['GPSDOP'])

            try:
                rev = requests.get(f'https://nominatim.openstreetmap.org/reverse', params={'format': 'json', 'lat': lat, 'lon': lon, 'addressdetails': 1}, headers={'User-Agent': 'Scanner/1.0'}, timeout=10)
                if rev.status_code == 200:
                    data = rev.json()
                    addr = data.get('address', {})
                    result['address'] = data.get('display_name', '')[:500]
                    result['country'] = addr.get('country', '')
                    result['city'] = addr.get('city') or addr.get('town') or addr.get('village', '')
                    result['street'] = addr.get('road', '')
                    result['postcode'] = addr.get('postcode', '')
            except:
                pass

            return result
        except Exception as e:
            print(f"[ERROR] GPS extraction failed: {e}")
            return None

class EXIFProcessor:
    @staticmethod
    def extract(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path}
            img = Image.open(file_path)
            result['image_width'] = img.width
            result['image_height'] = img.height

            exif = img._getexif()
            if exif:
                for k, v in exif.items():
                    tag = TAGS.get(k, str(k))
                    if isinstance(v, bytes):
                        try:
                            v = v.decode('utf-8', errors='ignore')
                        except:
                            v = str(v)

                    if tag == 'Make':
                        result['camera_make'] = str(v)
                    elif tag == 'Model':
                        result['camera_model'] = str(v)
                    elif tag == 'LensModel':
                        result['lens_model'] = str(v)
                    elif tag == 'FocalLength':
                        result['focal_length'] = str(v)
                    elif tag == 'FNumber':
                        result['aperture'] = str(v)
                    elif tag == 'ExposureTime':
                        result['shutter_speed'] = str(v)
                    elif tag == 'ISOSpeedRatings':
                        try:
                            result['iso'] = int(v)
                        except:
                            result['iso'] = 0
                    elif tag == 'DateTimeOriginal':
                        result['date_taken'] = str(v)
                    elif tag == 'Orientation':
                        try:
                            result['orientation'] = int(v)
                        except:
                            result['orientation'] = 1
            return result
        except Exception as e:
            print(f"[ERROR] EXIF extraction failed: {e}")
            return result

class FaceDetector:
    @staticmethod
    def detect(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path, 'face_count': 0, 'faces_json': '[]'}
            cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            img = cv2.imread(file_path)
            if img is None:
                return result
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            faces = cascade.detectMultiScale(gray, 1.1, 5)
            faces_list = [{'x': int(x), 'y': int(y), 'width': int(w), 'height': int(h)} for (x, y, w, h) in faces]
            result['face_count'] = len(faces)
            result['faces_json'] = json.dumps(faces_list)
            return result
        except Exception as e:
            print(f"[ERROR] Face detection failed: {e}")
            return result

class OCRProcessor:
    @staticmethod
    def extract(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path, 'extracted_text': '', 'word_count': 0, 'confidence': 0}
            try:
                import pytesseract
                img = cv2.imread(file_path)
                if img is None:
                    return result
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                text = pytesseract.image_to_string(thresh, lang='eng')
                result['extracted_text'] = text[:10000]
                result['word_count'] = len(text.split())
            except ImportError:
                pass
            return result
        except:
            return result

class StegoDetector:
    @staticmethod
    def detect(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path, 'has_hidden_data': 0, 'detection_method': '', 'confidence': 0}
            img = Image.open(file_path)
            arr = np.array(img)
            if len(arr.shape) == 3:
                bits = []
                for i in range(min(1000, arr.shape[0])):
                    for j in range(min(1000, arr.shape[1])):
                        for c in range(3):
                            bits.append(arr[i, j, c] & 1)
                bits_str = ''.join(str(b) for b in bits[:200])
                zeros = bits_str.count('0')
                ones = bits_str.count('1')
                total = zeros + ones
                if total > 0 and abs(zeros/total - 0.5) > 0.12:
                    result['has_hidden_data'] = 1
                    result['detection_method'] = 'lsb_imbalance'
                    result['confidence'] = 0.6
            return result
        except:
            return result

class DocumentParser:
    @staticmethod
    def parse(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path, 'page_count': 0, 'word_count': 0, 'character_count': 0, 'full_text': ''}
            text = ''
            if file_path.endswith('.pdf'):
                try:
                    from PyPDF2 import PdfReader
                    with open(file_path, 'rb') as f:
                        reader = PdfReader(f)
                        result['page_count'] = len(reader.pages)
                        if reader.metadata:
                            result['author'] = reader.metadata.get('/Author', '')
                            result['title'] = reader.metadata.get('/Title', '')
                            result['creator'] = reader.metadata.get('/Creator', '')
                        for page in reader.pages[:20]:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text
                except:
                    pass
            elif file_path.endswith('.docx'):
                try:
                    import docx
                    doc = docx.Document(file_path)
                    result['page_count'] = len(doc.paragraphs) // 30 + 1
                    text = '\n'.join([p.text for p in doc.paragraphs])
                    if doc.core_properties.author:
                        result['author'] = doc.core_properties.author
                    if doc.core_properties.title:
                        result['title'] = doc.core_properties.title
                except:
                    pass
            elif file_path.endswith('.txt'):
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read(20000)
                except:
                    pass
            result['full_text'] = text[:20000]
            result['word_count'] = len(text.split())
            result['character_count'] = len(text)
            return result
        except:
            return result

class AudioProcessor:
    @staticmethod
    def extract(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path}
            if file_path.endswith('.mp3'):
                try:
                    from mutagen.mp3 import MP3
                    audio = MP3(file_path)
                    result['duration_seconds'] = audio.info.length
                    result['bitrate'] = audio.info.bitrate
                    result['sample_rate'] = audio.info.sample_rate
                    mins = int(result['duration_seconds'] // 60)
                    secs = int(result['duration_seconds'] % 60)
                    result['duration_formatted'] = f"{mins}:{secs:02d}"
                    if audio.tags:
                        result['artist'] = str(audio.get('TPE1', [''])[0])
                        result['title'] = str(audio.get('TIT2', [''])[0])
                        result['album'] = str(audio.get('TALB', [''])[0])
                except:
                    pass
            elif file_path.endswith('.wav'):
                try:
                    import wave
                    with wave.open(file_path, 'rb') as wav:
                        frames = wav.getnframes()
                        rate = wav.getframerate()
                        result['duration_seconds'] = frames / rate
                        result['channels'] = wav.getnchannels()
                        result['sample_rate'] = rate
                        mins = int(result['duration_seconds'] // 60)
                        secs = int(result['duration_seconds'] % 60)
                        result['duration_formatted'] = f"{mins}:{secs:02d}"
                except:
                    pass
            return result
        except:
            return result

class VideoProcessor:
    @staticmethod
    def extract(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path}
            cap = cv2.VideoCapture(file_path)
            if cap.isOpened():
                result['width'] = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
                result['height'] = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
                result['fps'] = cap.get(cv2.CAP_PROP_FPS)
                result['frame_count'] = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
                if result['fps'] > 0:
                    result['duration_seconds'] = result['frame_count'] / result['fps']
                    mins = int(result['duration_seconds'] // 60)
                    secs = int(result['duration_seconds'] % 60)
                    result['duration_formatted'] = f"{mins}:{secs:02d}"
                cap.release()
            return result
        except:
            return result

class ArchiveProcessor:
    @staticmethod
    def analyze(file_path, scan_id):
        try:
            result = {'scan_id': scan_id, 'file_path': file_path, 'file_count': 0}
            if zipfile.is_zipfile(file_path):
                result['archive_type'] = 'zip'
                with zipfile.ZipFile(file_path, 'r') as zf:
                    files = zf.namelist()
                    result['file_count'] = len(files)
                    result['total_size'] = sum(zf.getinfo(f).file_size for f in files)
                    result['files_list'] = json.dumps(files[:100])
            elif tarfile.is_tarfile(file_path):
                result['archive_type'] = 'tar'
                with tarfile.open(file_path, 'r') as tf:
                    files = tf.getnames()
                    result['file_count'] = len(files)
                    result['files_list'] = json.dumps(files[:100])
            return result
        except:
            return result

class IPGeolocationService:
    def __init__(self):
        self.sources = [
            ('http://ip-api.com/json/', 'ip-api.com'),
            ('https://ipinfo.io/', 'ipinfo.io'),
            ('https://ipwhois.app/json/', 'ipwhois.app')
        ]

    def get_my_ip(self):
        try:
            r = requests.get('https://api.ipify.org?format=json', timeout=5)
            return r.json()['ip']
        except:
            return None

    def get_locations(self, scan_id, ip=None):
        locations = []
        for url, source in self.sources:
            try:
                full_url = f'{url}{ip}' if ip else url
                r = requests.get(full_url, timeout=5)
                d = r.json()
                if source == 'ip-api.com' and d.get('status') == 'success':
                    locations.append({
                        'scan_id': scan_id,
                        'ip_address': d.get('query'),
                        'latitude': d.get('lat'),
                        'longitude': d.get('lon'),
                        'city': d.get('city'),
                        'country': d.get('country'),
                        'region': d.get('regionName'),
                        'isp': d.get('isp'),
                        'organization': d.get('org'),
                        'asn': d.get('as'),
                        'timezone': d.get('timezone'),
                        'source_name': source
                    })
                elif source == 'ipinfo.io' and 'loc' in d:
                    lat, lon = d['loc'].split(',')
                    locations.append({
                        'scan_id': scan_id,
                        'ip_address': d.get('ip'),
                        'latitude': float(lat),
                        'longitude': float(lon),
                        'city': d.get('city'),
                        'country': d.get('country'),
                        'region': d.get('region'),
                        'isp': d.get('org'),
                        'organization': d.get('org'),
                        'timezone': d.get('timezone'),
                        'source_name': source
                    })
                elif source == 'ipwhois.app' and d.get('success') != False:
                    locations.append({
                        'scan_id': scan_id,
                        'ip_address': d.get('ip'),
                        'latitude': d.get('latitude'),
                        'longitude': d.get('longitude'),
                        'city': d.get('city'),
                        'country': d.get('country'),
                        'region': d.get('region'),
                        'isp': d.get('isp'),
                        'organization': d.get('org'),
                        'asn': d.get('asn'),
                        'timezone': d.get('timezone'),
                        'source_name': source
                    })
            except:
                continue
        return locations

class SystemInfoCollector:
    @staticmethod
    def collect(scan_id):
        try:
            result = {'scan_id': scan_id}
            result['hostname'] = socket.gethostname()
            result['os_name'] = platform.system() if hasattr(platform, 'system') else 'Unknown'
            result['os_version'] = platform.version() if hasattr(platform, 'version') else 'Unknown'
            result['python_version'] = sys.version[:50]
            try:
                import psutil
                result['cpu_cores'] = psutil.cpu_count()
                mem = psutil.virtual_memory()
                result['ram_total_gb'] = round(mem.total / (1024**3), 2)
                disk = psutil.disk_usage('/')
                result['disk_total_gb'] = round(disk.total / (1024**3), 2)
            except:
                result['cpu_cores'] = os.cpu_count()
                result['ram_total_gb'] = 0
                result['disk_total_gb'] = 0
            try:
                import cpuinfo
                cpu = cpuinfo.get_cpu_info()
                result['cpu_model'] = cpu.get('brand_raw', 'Unknown')
            except:
                result['cpu_model'] = 'Unknown'
            return result
        except:
            return {'scan_id': scan_id}

class NetworkInfoCollector:
    @staticmethod
    def collect(scan_id):
        try:
            result = {'scan_id': scan_id}
            try:
                import netifaces
                interfaces = netifaces.interfaces()
                for iface in interfaces:
                    addrs = netifaces.ifaddresses(iface)
                    if netifaces.AF_INET in addrs:
                        for addr in addrs[netifaces.AF_INET]:
                            if not addr['addr'].startswith('127.'):
                                result['interface_name'] = iface
                                result['ipv4'] = addr['addr']
                                break
                    if netifaces.AF_LINK in addrs and 'mac_address' not in result:
                        for addr in addrs[netifaces.AF_LINK]:
                            result['mac_address'] = addr.get('addr', '')
                            break
                try:
                    gateways = netifaces.gateways()
                    if 'default' in gateways and netifaces.AF_INET in gateways['default']:
                        result['gateway'] = gateways['default'][netifaces.AF_INET][0]
                except:
                    pass
            except:
                pass
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
                s.connect(('8.8.8.8', 80))
                result['external_ip'] = s.getsockname()[0]
                s.close()
            except:
                pass
            return result
        except:
            return {'scan_id': scan_id}

class MapGenerator:
    @staticmethod
    def generate(gps_data, scan_id=None):
        if not gps_data:
            return None
        if scan_id:
            path = os.path.join(GPS_PATH, f"map_{scan_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html")
        else:
            path = os.path.join(GPS_PATH, f"map_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html")
        center = [gps_data[0][1], gps_data[0][2]]
        m = folium.Map(location=center, zoom_start=12)
        cluster = MarkerCluster().add_to(m)
        for gps in gps_data:
            popup = f"<b>{os.path.basename(gps[0])}</b><br>{gps[3] or ''}<br>{gps[4]}, {gps[5]}"
            folium.Marker([gps[1], gps[2]], popup=popup).add_to(cluster)
        m.save(path)
        print(f"[MAP] Saved: {path}")
        return path

class JSONExporter:
    @staticmethod
    def export_scan(db, scan_id):
        if not db.scan_exists(scan_id):
            print(f"[ERROR] Scan {scan_id} not found in database")
            return None

        output_path = os.path.join(JSON_PATH, f"{scan_id}.json")
        result = {
            'export_time': datetime.now().isoformat(),
            'scan_id': scan_id,
            'scanner_version': '5.0'
        }

        cursor = db.conn.cursor()

        cursor.execute('SELECT * FROM scans WHERE scan_id=?', (scan_id,))
        row = cursor.fetchone()
        if row:
            result['scan_info'] = dict(row)
            result['scan_info']['json_export_path'] = output_path

        cursor.execute('SELECT * FROM files WHERE scan_id=?', (scan_id,))
        result['files'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM gps_data WHERE scan_id=?', (scan_id,))
        result['gps_locations'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM exif_data WHERE scan_id=?', (scan_id,))
        result['exif_data'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM documents WHERE scan_id=?', (scan_id,))
        result['documents'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM audio_data WHERE scan_id=?', (scan_id,))
        result['audio_files'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM video_data WHERE scan_id=?', (scan_id,))
        result['video_files'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM archives WHERE scan_id=?', (scan_id,))
        result['archives'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM faces_data WHERE scan_id=?', (scan_id,))
        result['faces'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM ocr_data WHERE scan_id=?', (scan_id,))
        result['ocr_texts'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM stego_data WHERE scan_id=?', (scan_id,))
        result['steganography'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM ip_data WHERE scan_id=?', (scan_id,))
        result['ip_geolocation'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM network_info WHERE scan_id=?', (scan_id,))
        result['network_info'] = [dict(r) for r in cursor.fetchall()]

        cursor.execute('SELECT * FROM system_info WHERE scan_id=?', (scan_id,))
        result['system_info'] = [dict(r) for r in cursor.fetchall()]

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False, default=str)

        cursor.execute('UPDATE scans SET json_export_path=? WHERE scan_id=?', (output_path, scan_id))
        db.conn.commit()

        print(f"[JSON] Exported to: {output_path}")
        print(f"[JSON] File size: {os.path.getsize(output_path)} bytes")
        return output_path

class UltimateScanner:
    def __init__(self):
        self.db = DatabaseManager()
        self.file_gatherer = FileInfoGatherer()
        self.ip_service = IPGeolocationService()

    def scan_single_file(self, file_path):
        scan_id = str(uuid.uuid4())[:8]
        start = datetime.now()

        if not os.path.exists(file_path):
            print(f"[ERROR] File not found: {file_path}")
            return None

        print(f"\n[SCAN] Starting scan of: {file_path}")
        print(f"[SCAN] Scan ID: {scan_id}")

        file_info = self.file_gatherer.get_info(file_path, scan_id)
        if file_info:
            self.db.save_file(file_info)
            mime = file_info.get('mime_type', '') or ''
            ext = file_info.get('extension', '') or ''
        else:
            mime = ''
            ext = os.path.splitext(file_path)[1].lower()

        gps = None
        doc = None
        audio = None
        video = None
        archive = None
        stego = None
        faces = None

        if mime.startswith('image/') or ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif', '.webp']:
            print("[PROCESS] Extracting GPS...")
            gps = GPSProcessor.extract(file_path, scan_id)
            if gps and gps.get('latitude'):
                self.db.save_gps(gps)
                print(f"[GPS] Found: {gps['latitude']:.6f}, {gps['longitude']:.6f}")

            print("[PROCESS] Extracting EXIF...")
            exif = EXIFProcessor.extract(file_path, scan_id)
            self.db.save_exif(exif)
            if exif.get('camera_make'):
                print(f"[EXIF] Camera: {exif.get('camera_make')} {exif.get('camera_model')}")

            print("[PROCESS] Detecting faces...")
            faces = FaceDetector.detect(file_path, scan_id)
            self.db.save_faces(faces)
            if faces['face_count'] > 0:
                print(f"[FACES] Found: {faces['face_count']}")

            print("[PROCESS] Running OCR...")
            ocr = OCRProcessor.extract(file_path, scan_id)
            self.db.save_ocr(ocr)
            if ocr['word_count'] > 0:
                print(f"[OCR] Extracted: {ocr['word_count']} words")

            print("[PROCESS] Checking steganography...")
            stego = StegoDetector.detect(file_path, scan_id)
            self.db.save_stego(stego)
            if stego['has_hidden_data']:
                print(f"[STEGO] Hidden data detected!")

        elif mime.startswith('text/') or ext in ['.pdf', '.docx', '.xlsx', '.txt']:
            print("[PROCESS] Parsing document...")
            doc = DocumentParser.parse(file_path, scan_id)
            self.db.save_document(doc)
            print(f"[DOC] Pages: {doc.get('page_count', 0)}, Words: {doc.get('word_count', 0)}")

        elif mime.startswith('audio/'):
            print("[PROCESS] Extracting audio metadata...")
            audio = AudioProcessor.extract(file_path, scan_id)
            self.db.save_audio(audio)
            if audio.get('duration_formatted'):
                print(f"[AUDIO] Duration: {audio['duration_formatted']}")

        elif mime.startswith('video/'):
            print("[PROCESS] Extracting video metadata...")
            video = VideoProcessor.extract(file_path, scan_id)
            self.db.save_video(video)
            if video.get('width'):
                print(f"[VIDEO] Resolution: {video['width']}x{video['height']}")

        elif ext in ['.zip', '.tar', '.gz']:
            print("[PROCESS] Analyzing archive...")
            archive = ArchiveProcessor.analyze(file_path, scan_id)
            self.db.save_archive(archive)
            print(f"[ARCHIVE] Files: {archive.get('file_count', 0)}")

        print("[PROCESS] Saving to database...")
        end = datetime.now()
        duration = (end - start).total_seconds()

        scan_data = {
            'scan_id': scan_id,
            'start_time': start.isoformat(),
            'end_time': end.isoformat(),
            'scan_path': file_path,
            'total_files': 1,
            'total_gps': 1 if gps and gps.get('latitude') else 0,
            'total_documents': 1 if doc and doc.get('word_count', 0) > 0 else 0,
            'total_audio': 1 if audio and audio.get('duration_seconds', 0) > 0 else 0,
            'total_video': 1 if video and video.get('width', 0) > 0 else 0,
            'total_archives': 1 if archive and archive.get('file_count', 0) > 0 else 0,
            'total_stego': 1 if stego and stego['has_hidden_data'] else 0,
            'total_faces': faces['face_count'] if faces else 0,
            'duration_seconds': duration,
            'scan_type': 'file',
            'computer_name': socket.gethostname(),
            'json_export_path': None
        }
        self.db.save_scan(scan_data)

        print("[PROCESS] Exporting to JSON...")
        json_path = JSONExporter.export_scan(self.db, scan_id)

        if gps and gps.get('latitude'):
            print("[PROCESS] Generating GPS map...")
            gps_data = self.db.get_gps_by_scan(scan_id)
            if gps_data:
                map_path = MapGenerator.generate(gps_data, scan_id)
                print(f"[MAP] Saved: {map_path}")

        print(f"\n[SUCCESS] Scan completed!")
        print(f"[INFO] Scan ID: {scan_id}")
        print(f"[INFO] Duration: {duration:.2f} seconds")
        print(f"[INFO] Database: {DB_PATH}")
        print(f"[INFO] JSON export: {json_path}")

        return scan_id

    def scan_folder(self, folder_path):
        scan_id = str(uuid.uuid4())[:8]
        start = datetime.now()

        if not os.path.exists(folder_path):
            print(f"[ERROR] Folder not found: {folder_path}")
            return None

        all_files = []
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                all_files.append(os.path.join(root, file))

        total = len(all_files)
        print(f"\n[SCAN] Starting folder scan: {folder_path}")
        print(f"[SCAN] Scan ID: {scan_id}")
        print(f"[SCAN] Total files found: {total}")

        gps_count = 0
        doc_count = 0
        audio_count = 0
        video_count = 0
        archive_count = 0
        stego_count = 0
        total_faces = 0

        for i, file_path in enumerate(all_files):
            print(f"\n[{i+1}/{total}] Processing: {os.path.basename(file_path)}")

            file_info = self.file_gatherer.get_info(file_path, scan_id)
            if file_info:
                self.db.save_file(file_info)
                mime = file_info.get('mime_type', '') or ''
                ext = file_info.get('extension', '') or ''
            else:
                mime = ''
                ext = os.path.splitext(file_path)[1].lower()

            if mime.startswith('image/') or ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif', '.webp']:
                gps = GPSProcessor.extract(file_path, scan_id)
                if gps and gps.get('latitude'):
                    self.db.save_gps(gps)
                    gps_count += 1
                    print(f"  ✓ GPS found")

                exif = EXIFProcessor.extract(file_path, scan_id)
                self.db.save_exif(exif)

                faces = FaceDetector.detect(file_path, scan_id)
                self.db.save_faces(faces)
                if faces['face_count'] > 0:
                    total_faces += faces['face_count']

                ocr = OCRProcessor.extract(file_path, scan_id)
                self.db.save_ocr(ocr)

                stego = StegoDetector.detect(file_path, scan_id)
                self.db.save_stego(stego)
                if stego['has_hidden_data']:
                    stego_count += 1
                    print(f"  ⚠ Hidden data detected")

            elif mime.startswith('text/') or ext in ['.pdf', '.docx', '.xlsx', '.txt']:
                doc = DocumentParser.parse(file_path, scan_id)
                self.db.save_document(doc)
                if doc.get('word_count', 0) > 0:
                    doc_count += 1
                    print(f"  ✓ Document: {doc['word_count']} words")

            elif mime.startswith('audio/'):
                audio = AudioProcessor.extract(file_path, scan_id)
                self.db.save_audio(audio)
                if audio.get('duration_seconds', 0) > 0:
                    audio_count += 1
                    print(f"  ✓ Audio: {audio.get('duration_formatted', 'unknown')}")

            elif mime.startswith('video/'):
                video = VideoProcessor.extract(file_path, scan_id)
                self.db.save_video(video)
                if video.get('width', 0) > 0:
                    video_count += 1
                    print(f"  ✓ Video: {video['width']}x{video['height']}")

            elif ext in ['.zip', '.tar', '.gz']:
                archive = ArchiveProcessor.analyze(file_path, scan_id)
                self.db.save_archive(archive)
                if archive.get('file_count', 0) > 0:
                    archive_count += 1
                    print(f"  ✓ Archive: {archive['file_count']} files")

        print("\n[PROCESS] Collecting IP geolocation...")
        ip_locations = self.ip_service.get_locations(scan_id)
        for ip_loc in ip_locations:
            self.db.save_ip(ip_loc)
        print(f"  Found {len(ip_locations)} IP location sources")

        print("[PROCESS] Collecting system info...")
        sys_info = SystemInfoCollector.collect(scan_id)
        self.db.save_system(sys_info)

        print("[PROCESS] Collecting network info...")
        net_info = NetworkInfoCollector.collect(scan_id)
        self.db.save_network(net_info)

        end = datetime.now()
        duration = (end - start).total_seconds()

        scan_data = {
            'scan_id': scan_id,
            'start_time': start.isoformat(),
            'end_time': end.isoformat(),
            'scan_path': folder_path,
            'total_files': total,
            'total_gps': gps_count,
            'total_documents': doc_count,
            'total_audio': audio_count,
            'total_video': video_count,
            'total_archives': archive_count,
            'total_stego': stego_count,
            'total_faces': total_faces,
            'duration_seconds': duration,
            'scan_type': 'folder',
            'computer_name': socket.gethostname(),
            'json_export_path': None
        }
        self.db.save_scan(scan_data)

        print("[PROCESS] Exporting to JSON...")
        json_path = JSONExporter.export_scan(self.db, scan_id)

        print("[PROCESS] Generating GPS map...")
        gps_data = self.db.get_gps_by_scan(scan_id)
        if gps_data:
            map_path = MapGenerator.generate(gps_data, scan_id)
            print(f"[MAP] Saved: {map_path}")

        print(f"\n[SUCCESS] Folder scan completed!")
        print(f"[INFO] Scan ID: {scan_id}")
        print(f"[INFO] Duration: {duration:.2f} seconds")
        print(f"[INFO] Total files: {total}")
        print(f"[INFO] GPS locations: {gps_count}")
        print(f"[INFO] Faces detected: {total_faces}")
        print(f"[INFO] Hidden data: {stego_count}")
        print(f"[INFO] Documents: {doc_count}, Audio: {audio_count}, Video: {video_count}, Archives: {archive_count}")
        print(f"[INFO] Database: {DB_PATH}")
        print(f"[INFO] JSON export: {json_path}")

        return scan_id

    def show_report(self, scan_id):
        print(f"\n[REPORT] Scan ID: {scan_id}")
        print("="*60)

        if not self.db.scan_exists(scan_id):
            print(f"[ERROR] Scan {scan_id} not found")
            return

        scans = self.db.get_scans()
        scan_data = None
        for s in scans:
            if s[0] == scan_id:
                scan_data = s
                break

        if scan_data:
            print(f"Date: {scan_data[1][:19] if scan_data[1] else 'N/A'}")
            print(f"Path: {scan_data[3]}")
            print(f"Total files: {scan_data[4]}")
            print(f"GPS locations: {scan_data[5]}")
            print(f"Faces detected: {scan_data[11]}")
            print(f"Steganography: {scan_data[10]}")
            print(f"Documents: {scan_data[6]}")
            print(f"Audio files: {scan_data[7]}")
            print(f"Video files: {scan_data[8]}")
            print(f"Archives: {scan_data[9]}")
            print(f"Duration: {scan_data[12]:.2f}s")
            print(f"JSON export: {scan_data[15]}")

        gps_list = self.db.get_gps_by_scan(scan_id)
        if gps_list:
            print(f"\n[GPS] Locations found: {len(gps_list)}")
            for g in gps_list[:10]:
                print(f"  {g[2]}: {g[3]:.6f}, {g[4]:.6f} - {g[7]}")

    def show_map(self, scan_id=None):
        if scan_id:
            gps_data = self.db.get_gps_by_scan(scan_id)
            if not gps_data:
                print(f"[ERROR] No GPS data for scan {scan_id}")
                return
            path = MapGenerator.generate(gps_data, scan_id)
        else:
            gps_data = self.db.get_all_gps()
            if not gps_data:
                print("[ERROR] No GPS data found in database")
                return
            path = MapGenerator.generate(gps_data)

        if path:
            webbrowser.open(f'file://{os.path.abspath(path)}')
            print(f"[MAP] Opened: {path}")

    def show_history(self):
        scans = self.db.get_scans()
        if not scans:
            print("[INFO] No scan history found")
            return

        print("\n[HISTORY] Scan History")
        print("="*80)
        for s in scans:
            print(f"\nID: {s[0]}")
            print(f"  Date: {s[1][:19] if s[1] else 'N/A'}")
            print(f"  Path: {s[3][:60]}")
            print(f"  Files: {s[4]} | GPS: {s[5]} | Faces: {s[11]} | Stego: {s[10]}")
            print(f"  JSON: {s[15]}")

    def export_json(self, scan_id=None):
        if scan_id:
            if not self.db.scan_exists(scan_id):
                print(f"[ERROR] Scan {scan_id} not found")
                return
            path = JSONExporter.export_scan(self.db, scan_id)
            print(f"[JSON] Exported: {path}")
        else:
            scans = self.db.get_scans()
            if not scans:
                print("[ERROR] No scans found")
                return
            print(f"[JSON] Exporting all {len(scans)} scans...")
            for s in scans:
                path = JSONExporter.export_scan(self.db, s[0])
                print(f"  {s[0]}: {path}")

    def show_ip(self):
        print("\n[IP] Geolocation Information")
        print("="*50)
        ip = self.ip_service.get_my_ip()
        print(f"Your IP address: {ip}")

        locs = self.ip_service.get_locations('temp')
        for loc in locs:
            print(f"\nSource: {loc.get('source_name')}")
            print(f"  Location: {loc.get('city')}, {loc.get('country')}")
            print(f"  Coordinates: {loc.get('latitude')}, {loc.get('longitude')}")
            print(f"  ISP: {loc.get('isp', 'N/A')}")
            print(f"  Timezone: {loc.get('timezone', 'N/A')}")

    def run(self):
        print("\n" + "="*70)
        print("ULTIMATE SCANNER V5.0 - FULL METADATA EXTRACTOR")
        print("="*70)

        while True:
            print("\n" + "="*70)
            print("MAIN MENU")
            print("="*70)
            print("1. Scan Single File")
            print("2. Scan Folder (Recursive)")
            print("3. Show Scan Report")
            print("4. Show GPS Map")
            print("5. Show Scan History")
            print("6. Export to JSON")
            print("7. IP Geolocation Info")
            print("8. Exit")
            print("="*70)

            choice = input("\nChoice: ").strip()

            if choice == '1':
                path = input("File path: ").strip('"')
                if os.path.exists(path):
                    self.scan_single_file(path)
                else:
                    print(f"[ERROR] File not found: {path}")

            elif choice == '2':
                path = input("Folder path: ").strip('"')
                if os.path.exists(path):
                    self.scan_folder(path)
                else:
                    print(f"[ERROR] Folder not found: {path}")

            elif choice == '3':
                sid = input("Scan ID: ").strip()
                self.show_report(sid)

            elif choice == '4':
                sid = input("Scan ID (or press Enter for all): ").strip()
                if sid:
                    self.show_map(sid)
                else:
                    self.show_map()

            elif choice == '5':
                self.show_history()

            elif choice == '6':
                sid = input("Scan ID (or press Enter for all): ").strip()
                if sid:
                    self.export_json(sid)
                else:
                    self.export_json()

            elif choice == '7':
                self.show_ip()

            elif choice == '8':
                print("\n[INFO] Closing database...")
                self.db.close()
                print("[INFO] Goodbye!")
                break

            else:
                print("[ERROR] Invalid choice")

if __name__ == "__main__":
    scanner = UltimateScanner()
    scanner.run()
